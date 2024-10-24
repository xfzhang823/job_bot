"""May not need this one"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pydantic import BaseModel, Field
from typing import List, Dict, Optional, Tuple
from enum import Enum
import re
from pathlib import Path


class StyleType(Enum):
    HEADING = "heading"
    BULLET = "bullet"
    NORMAL = "normal"
    CONTACT = "contact"
    DATE = "date"


class StyleProperties(BaseModel):
    name: str
    font_name: Optional[str]
    font_size: Optional[int]
    is_bold: bool = False
    is_italic: bool = False
    alignment: str
    left_indent: float = 0.0
    space_before: float = 0.0
    space_after: float = 0.0


class SectionTemplate(BaseModel):
    name: str
    start_index: int
    end_index: Optional[int]
    style: StyleProperties
    subsection_styles: List[StyleProperties] = Field(default_factory=list)
    bullet_style: Optional[StyleProperties] = None


class DocumentTemplate(BaseModel):
    sections: List[SectionTemplate]
    global_styles: Dict[StyleType, StyleProperties]
    section_order: List[str]
    indentation_levels: Dict[int, float]


class TemplateExtractor:
    def __init__(self):
        self.common_section_headers = {
            "experience": [
                "experience",
                "work experience",
                "professional experience",
                "employment history",
            ],
            "education": ["education", "educational background", "academic background"],
            "skills": ["skills", "technical skills", "core competencies"],
            "summary": ["summary", "professional summary", "profile"],
            "contact": ["contact", "contact information", "personal information"],
        }

    def _get_style_properties(self, paragraph: Paragraph) -> StyleProperties:
        """Extract detailed style properties from a paragraph"""
        style = paragraph.style
        font = style.font
        paragraph_format = style.paragraph_format

        return StyleProperties(
            name=style.name,
            font_name=font.name if font.name else None,
            font_size=font.size.pt if font.size else None,
            is_bold=font.bold if font.bold else False,
            is_italic=font.italic if font.italic else False,
            alignment=str(paragraph.alignment) if paragraph.alignment else "LEFT",
            left_indent=(
                paragraph_format.left_indent.pt if paragraph_format.left_indent else 0.0
            ),
            space_before=(
                paragraph_format.space_before.pt
                if paragraph_format.space_before
                else 0.0
            ),
            space_after=(
                paragraph_format.space_after.pt if paragraph_format.space_after else 0.0
            ),
        )

    def _identify_section_type(self, text: str) -> Optional[str]:
        """Identify the type of section based on its header text"""
        text = text.lower().strip()
        for section_type, headers in self.common_section_headers.items():
            if text in headers:
                return section_type
        return None

    def _is_bullet_point(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is a bullet point"""
        return bool(paragraph.style.paragraph_format.left_indent) and (
            paragraph.text.strip().startswith("•")
            or paragraph.text.strip().startswith("-")
            or bool(re.match(r"^\s*[\u2022\-\*]\s", paragraph.text))
        )

    def _is_date_format(self, text: str) -> bool:
        """Check if text contains a date pattern"""
        date_patterns = [
            r"\d{1,2}/\d{1,2}/\d{2,4}",
            r"\d{1,2}-\d{1,2}-\d{2,4}",
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}",
            r"Present|Current|Now",
        ]
        return any(re.search(pattern, text) for pattern in date_patterns)

    def extract_template(self, doc_path: Path) -> DocumentTemplate:
        """
        Extract complete template information from a document
        Returns a DocumentTemplate with all structure and styling information
        """
        doc = Document(doc_path)
        sections: List[SectionTemplate] = []
        current_section: Optional[SectionTemplate] = None
        global_styles: Dict[StyleType, StyleProperties] = {}
        indentation_levels: Dict[int, float] = {}

        for idx, para in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                continue

            style_props = self._get_style_properties(para)

            # Track indentation levels
            if style_props.left_indent > 0:
                indentation_levels[idx] = style_props.left_indent

            # Identify section headers
            if para.style.name.startswith("Heading") or self._identify_section_type(
                para.text
            ):
                # Close previous section
                if current_section:
                    current_section.end_index = idx - 1
                    sections.append(current_section)

                # Start new section
                current_section = SectionTemplate(
                    name=para.text.strip(),
                    start_index=idx,
                    end_index=None,
                    style=style_props,
                )

            # Process subsection styles and bullet points
            elif current_section:
                if self._is_bullet_point(para):
                    if not current_section.bullet_style:
                        current_section.bullet_style = style_props
                else:
                    current_section.subsection_styles.append(style_props)

            # Collect global styles
            if self._is_bullet_point(para):
                global_styles[StyleType.BULLET] = style_props
            elif para.style.name.startswith("Heading"):
                global_styles[StyleType.HEADING] = style_props
            elif self._is_date_format(para.text):
                global_styles[StyleType.DATE] = style_props
            else:
                global_styles[StyleType.NORMAL] = style_props

        # Close last section
        if current_section:
            current_section.end_index = len(doc.paragraphs) - 1
            sections.append(current_section)

        return DocumentTemplate(
            sections=sections,
            global_styles=global_styles,
            section_order=[section.name for section in sections],
            indentation_levels=indentation_levels,
        )

    def analyze_template(self, template: DocumentTemplate) -> Dict:
        """
        Analyze the template and return useful metadata
        """
        return {
            "section_count": len(template.sections),
            "indentation_levels": len(template.indentation_levels),
            "has_bullet_points": StyleType.BULLET in template.global_styles,
            "section_structure": [
                {
                    "name": section.name,
                    "length": section.end_index - section.start_index + 1,
                    "has_subsections": len(section.subsection_styles) > 0,
                    "has_bullets": section.bullet_style is not None,
                }
                for section in template.sections
            ],
        }


# Example usage with a sample resume
def example_template_extraction():
    """
    Example of how to use the template extractor
    """
    # Create a sample resume
    doc = Document()

    # Add contact information
    doc.add_heading("John Doe", 0)
    contact = doc.add_paragraph(
        "john.doe@email.com | (555) 555-5555 | LinkedIn: johndoe"
    )
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Experience section
    doc.add_heading("Professional Experience", 1)
    doc.add_paragraph("Senior Software Engineer", style="Heading 2")
    doc.add_paragraph("Tech Corp | Jan 2020 - Present")
    bullet1 = doc.add_paragraph("• Led development of cloud-based application")
    bullet1.style.paragraph_format.left_indent = Pt(36)
    bullet2 = doc.add_paragraph("• Improved system performance by 40%")
    bullet2.style.paragraph_format.left_indent = Pt(36)

    # Add Education section
    doc.add_heading("Education", 1)
    doc.add_paragraph("Master of Computer Science", style="Heading 2")
    doc.add_paragraph("University of Technology | 2018 - 2020")

    # Save the sample document
    sample_path = Path("sample_resume.docx")
    doc.save(sample_path)

    # Extract and analyze template
    extractor = TemplateExtractor()
    template = extractor.extract_template(sample_path)
    analysis = extractor.analyze_template(template)

    print("Template Analysis:", analysis)

    return template, analysis


if __name__ == "__main__":
    template, analysis = example_template_extraction()
