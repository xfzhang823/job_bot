from pathlib import Path
from datetime import date
from typing import List, Optional, Dict, Tuple
from pydantic import BaseModel, Field
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
import json


# Pydantic Models
class DocumentLocation(BaseModel):
    """Tracks location in Word document for reconstruction"""

    paragraph_index: int
    is_heading: bool
    level: Optional[int] = None
    style_name: str


class ContentBlock(BaseModel):
    """Maps content to its location in document"""

    content: str
    location: DocumentLocation
    identifier: str


class ExperienceItem(BaseModel):
    company: str
    title: str
    start_date: date
    end_date: Optional[date]
    bullets: List[str]
    location: Optional[str]
    _doc_locations: Dict[str, DocumentLocation] = Field(default_factory=dict)


class EducationItem(BaseModel):
    institution: str
    degree: str
    graduation_date: date
    gpa: Optional[float]
    _doc_locations: Dict[str, DocumentLocation] = Field(default_factory=dict)


class ResumeSchema(BaseModel):
    personal_info: Dict[str, str]
    experience: List[ExperienceItem]
    education: List[EducationItem]
    skills: List[str]
    _template_metadata: Dict[str, str] = Field(default_factory=dict)


class ResumePipeline:
    def __init__(self):
        self.template_fingerprint: Dict = {}
        self.content_mappings: Dict[str, ContentBlock] = {}

    def extract_template(self, doc: DocxDocument) -> None:
        """Analyze and store document template structure"""
        self.template_fingerprint = {
            "styles": [p.style.name for p in doc.paragraphs],
            "levels": [
                p.style.base_style.name if p.style.base_style else None
                for p in doc.paragraphs
            ],
            "section_markers": self._identify_sections(doc),
        }

    def _identify_sections(self, doc: DocxDocument) -> Dict[str, DocumentLocation]:
        """Identify major resume sections and their locations"""
        sections = {}
        for idx, para in enumerate(doc.paragraphs):
            if para.style.name.startswith("Heading"):
                sections[para.text.strip().lower()] = DocumentLocation(
                    paragraph_index=idx,
                    is_heading=True,
                    level=int(para.style.name[-1]),
                    style_name=para.style.name,
                )
        return sections

    def parse_document(self, doc_path: Path) -> ResumeSchema:
        """Parse Word document into structured data"""
        doc = Document(doc_path)
        self.extract_template(doc)

        # Initialize resume data
        resume_data = {
            "personal_info": self._extract_personal_info(doc),
            "experience": self._extract_experience(doc),
            "education": self._extract_education(doc),
            "skills": self._extract_skills(doc),
        }

        # Validate through Pydantic
        return ResumeSchema(**resume_data)

    def _extract_section_content(
        self, doc: DocxDocument, start_marker: str, end_marker: str
    ) -> List[Paragraph]:
        """Extract content between two section markers"""
        content = []
        capturing = False

        for para in doc.paragraphs:
            if para.text.strip().lower() == start_marker.lower():
                capturing = True
                continue
            elif para.text.strip().lower() == end_marker.lower():
                break

            if capturing and para.text.strip():
                content.append(para)

        return content

    def _extract_experience(self, doc: DocxDocument) -> List[Dict]:
        """Extract work experience entries"""
        experience_section = self._extract_section_content(
            doc, "Experience", "Education"
        )

        experiences = []
        current_experience = {}

        for para in experience_section:
            # Logic to parse experience entries
            # Store document locations for reconstruction
            if para.style.name.startswith("Heading"):
                if current_experience:
                    experiences.append(current_experience)
                current_experience = {
                    "_doc_locations": {
                        "title": DocumentLocation(
                            paragraph_index=para._p.index,
                            is_heading=True,
                            level=int(para.style.name[-1]),
                            style_name=para.style.name,
                        )
                    }
                }
            # Add more parsing logic here

        return experiences

    def update_resume(
        self, resume: ResumeSchema, updates: Dict[str, any]
    ) -> ResumeSchema:
        """Update resume content while maintaining structure"""
        updated_data = resume.dict()

        for key, value in updates.items():
            if "." in key:  # Handle nested updates
                section, field = key.split(".", 1)
                if section in updated_data:
                    if isinstance(updated_data[section], list):
                        # Handle list updates
                        pass
                    else:
                        # Handle direct field updates
                        updated_data[section][field] = value
            else:
                updated_data[key] = value

        return ResumeSchema(**updated_data)

    def reconstruct_document(
        self, resume: ResumeSchema, template_doc: Path
    ) -> Document:
        """Reconstruct Word document from updated data"""
        doc = Document(template_doc)

        # Reconstruct using template fingerprint
        for section, location in self.template_fingerprint["section_markers"].items():
            if section in resume.dict():
                self._rebuild_section(doc, section, resume.dict()[section], location)

        return doc

    def _rebuild_section(
        self, doc: Document, section_name: str, content: any, location: DocumentLocation
    ) -> None:
        """Rebuild a section while maintaining formatting"""
        # Implementation for rebuilding document sections
        pass


# Usage Example
def process_resume(doc_path: str, updates: Dict[str, any]) -> None:
    pipeline = ResumePipeline()

    # Parse original document
    resume = pipeline.parse_document(Path(doc_path))

    # Apply updates
    updated_resume = pipeline.update_resume(resume, updates)

    # Reconstruct document
    new_doc = pipeline.reconstruct_document(updated_resume, Path(doc_path))

    # Save updated document
    new_doc.save("updated_resume.docx")
