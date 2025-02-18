"""
TBA
"""

import logging
import logging_config

from pydantic import BaseModel  # For structured data handling (optional)
from typing import Dict, Optional, Union  # For type annotations in functions
from pathlib import Path  # To handle file paths
import json  # For converting between dict and JSON format

from docx import Document  # Main library for reading and writing DOCX files
from docx.shared import Pt  # To set indentation in points
from docx.oxml.ns import qn  # For working with bullet/numbering formatting in DOCX
from docx.oxml import OxmlElement  # To directly manipulate XML for list styles
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import the alignment enum

from utils.generic_utils import save_to_json_file


logger = logging.getLogger(__name__)


class DocxJsonProcessor:
    def __init__(self, docx_path: Union[Path, str]):
        self.doc = Document(str(docx_path))
        self.logger = logging.getLogger(__name__)

    def add_hyperlink(self, paragraph, url, text):
        """
        Adds a hyperlink to a paragraph in the docx.
        :param paragraph: The paragraph we want to add the hyperlink to.
        :param url: The URL to link to.
        :param text: The display text for the link.
        :return: None
        """
        # Create the w:hyperlink tag and add needed values
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create a run element (w:r)
        run = OxmlElement("w:r")

        # Create run properties for the hyperlink (underline, blue color)
        rPr = OxmlElement("w:rPr")

        # Underline element
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")  # Single underline
        rPr.append(u)

        # Color element (typically blue for hyperlinks)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")  # Blue color
        rPr.append(color)

        # Append the run properties to the run
        run.append(rPr)

        # Create a text element (w:t)
        text_elem = OxmlElement("w:t")

        # Set the text content for the text element
        text_elem._set_text(text)  # Use `_set_text` instead of `.text`

        # Append the text element to the run element
        run.append(text_elem)

        # Append the run element to the hyperlink element
        hyperlink.append(run)

        # Finally, append the hyperlink element to the paragraph
        paragraph._p.append(hyperlink)

    def extract_to_json(self) -> Dict:
        """Extract DOCX content into a structured JSON format with indents and bullets."""
        content = {}
        section_index = -1  # Start at -1 to increment when a heading is found

        self.logger.info("Starting DOCX to JSON extraction...")

        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                self.logger.debug(f"Skipping empty paragraph at index {i}")
                continue  # Skip empty paragraphs

            # Extract paragraph information
            info = self._get_paragraph_info(para)

            self.logger.debug(f"Processing paragraph {i}: {info['text'][:50]}...")
            self.logger.debug(
                f"Style: {info['style_name']}, Indent: {info['indent']}, Is bullet: {info['is_bullet']}"
            )

            if info["style_name"].startswith("Heading"):
                # This marks a new section (heading)
                section_index += 1
                self.logger.info(
                    f"New section detected: {info['text'][:50]} (Section Index: {section_index})"
                )
                content[f"{section_index}.responsibilities"] = (
                    []
                )  # Create a new section
            else:
                # Create a default section if no heading has been found yet
                if section_index == -1:
                    self.logger.info(
                        "No heading found. Using default section for unassigned paragraphs."
                    )
                    section_index = 0  # Use 0 as the default section index
                    content[f"{section_index}.responsibilities"] = []

                # Add paragraph content with indent and bullet info
                content[f"{section_index}.responsibilities"].append(
                    {
                        "text": info["text"],
                        "indent": info["indent"],
                        "is_bullet": info["is_bullet"],
                    }
                )

                self.logger.debug(
                    f"Appended responsibility to section {section_index}: {info['text'][:50]}..."
                )

        self.logger.info("Finished extracting DOCX content to JSON.")
        self.logger.debug(f"Extracted content: {content}")

        return self.flatten_json(content)

    def add_bullet_paragraph(self, doc, text, indent):
        """Add a bulleted paragraph with specified indentation."""
        p = doc.add_paragraph(text)
        p_format = p.paragraph_format
        p_format.left_indent = Pt(indent)
        p.style = "List Bullet"  # Assign bullet list style

    def flatten_json(self, nested_dict: Dict) -> Dict:
        """Flatten a nested dictionary."""
        flat_dict = {}

        def _flatten(obj, key_prefix=""):
            if isinstance(obj, dict):
                for k, v in obj.items():
                    _flatten(v, f"{key_prefix}{k}.")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    _flatten(item, f"{key_prefix}{i}.")
            else:
                flat_dict[key_prefix[:-1]] = obj  # Remove last period

        _flatten(nested_dict)
        return flat_dict

    def json_to_docx(self, json_data: Dict, output_docx_path: str):
        """Convert JSON data back to DOCX format with bullets, alignment, bold, underline,
        and hyperlinks."""
        self.logger.info("Starting JSON to DOCX conversion...")
        doc = Document()
        section_index = -1  # Track the section being processed

        current_para_info = {}  # Store info for the current paragraph

        for key, para_info in json_data.items():
            self.logger.debug(f"Processing key: {key}, para_info: {para_info}")

            # Extract paragraph identifier (e.g., 0.responsibilities.0)
            # and the attribute (e.g., text, indent)
            key_parts = key.split(".")
            para_attr = key_parts[-1]  # This will be 'text', 'indent', or 'is_bullet'
            para_id = ".".join(key_parts[:-1])  # This will be '0.responsibilities.0'
            section_num = int(key_parts[0])  # Extract section number

            # If this is a new paragraph ID, process the previously stored paragraph
            if para_id != current_para_info.get("id", ""):
                if "id" in current_para_info:
                    # Add the previously processed paragraph
                    self._add_paragraph(doc, current_para_info)
                # Reset for the new paragraph
                current_para_info = {"id": para_id}

            # Store the attribute in current_para_info
            current_para_info[para_attr] = para_info

            # If a new section is detected, add a heading
            if section_num != section_index:
                section_index = section_num
                doc.add_heading(f"Section {section_num}", level=1)
                self.logger.info(f"Adding new section: Section {section_num}")

        # Handle the last paragraph after the loop
        if "id" in current_para_info:
            self._add_paragraph(doc, current_para_info)

        # Save the document
        doc.save(output_docx_path)
        self.logger.info(
            f"JSON to DOCX conversion complete. Saved to {output_docx_path}"
        )

    def _add_paragraph(self, doc, para_info):
        """Helper function to add a paragraph to the document based on the info."""
        para_text = para_info.get("text", "")
        indent = para_info.get("indent", 0)
        is_bullet = para_info.get("is_bullet", False)
        alignment = para_info.get("alignment", "left")
        is_bold = para_info.get("is_bold", False)
        is_underline = para_info.get("is_underline", False)
        hyperlinks = para_info.get("hyperlinks", [])

        # Add the paragraph with the extracted info
        p = doc.add_paragraph()  # Create empty paragraph
        p.paragraph_format.left_indent = Pt(indent)

        # If there are hyperlinks, split text into normal and hyperlink parts
        if hyperlinks:
            start_index = 0
            for hyperlink in hyperlinks:
                url = hyperlink.get("url", "")
                link_text = hyperlink.get("text", "")

                # Add non-hyperlinked text before the hyperlink (if any)
                if para_text[start_index : start_index + len(link_text)]:
                    p.add_run(para_text[start_index : start_index + len(link_text)])

                # Add hyperlink text
                self.add_hyperlink(p, url, link_text)
                start_index += len(link_text)

            # Add remaining non-hyperlinked text
            if start_index < len(para_text):
                p.add_run(para_text[start_index:])
        else:
            # Add plain text without hyperlinks
            run = p.add_run(para_text)

        # Add run (to apply bold/underline inside the run)
        run = p.add_run(para_text)
        if is_bold:
            run.bold = True
        if is_underline:
            run.underline = True

        # Set alignment using WD_ALIGN_PARAGRAPH
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Default to left if none is provided

        # Apply bullet style if needed
        if is_bullet:
            p.style = "List Bullet"

        # Add hyperlinks to the paragraph
        for link in hyperlinks:
            self.add_hyperlink(p, link["url"], link["text"])

    def _get_paragraph_info(self, para):
        """Extract detailed paragraph information from a DOCX paragraph."""
        para_text = para.text.strip()  # Get paragraph text
        indent = (
            para.paragraph_format.left_indent.pt
            if para.paragraph_format.left_indent
            else 0
        )
        is_bullet = para.style.name.lower().startswith("list bullet")

        # Extract alignment
        alignment = "left"
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            alignment = "center"
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            alignment = "right"
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            alignment = "justify"

        # Extract bold, underline, and hyperlinks from runs
        is_bold = False
        is_underline = False
        hyperlinks = []

        for run in para.runs:
            if run.bold:
                is_bold = True
            if run.underline:
                is_underline = True

            # Check for hyperlinks in the run XML
            r = run._r  # Get the underlying XML element for the run
            hyperlink = r.find(qn("w:hyperlink"))  # Search for hyperlink in the XML
            if hyperlink is not None:
                r_id = hyperlink.get(qn("r:id"))
                if r_id:
                    # Resolve the actual URL using the paragraph's relationship data
                    part = para.part  # Part that contains the relationship data
                    rel = part.rels[r_id]  # Find the relationship by id
                    url = rel.target_ref  # Get the actual URL
                    hyperlinks.append({"url": url, "text": run.text})

        return {
            "text": para_text,
            "indent": indent,
            "is_bullet": is_bullet,
            "alignment": alignment,
            "is_bold": is_bold,
            "is_underline": is_underline,
            "hyperlinks": hyperlinks,
        }


# Example Usage:
if __name__ == "__main__":
    doc_path = Path("your_resume.docx")  # Replace with your actual DOCX file path
    processor = DocxJsonProcessor(doc_path)

    # Convert DOCX to JSON
    json_content = processor.extract_to_json()

    # Save JSON content to a file (optional)
    save_to_json_file(json_content, "file_path")
